from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Color palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x37, 0x6C)   # deep navy
ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)

FONT_BODY    = "Calibri"
FONT_HEADING = "Calibri Light"

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Fill a table cell with a solid background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_para_border_bottom(para, color_hex="2E74B5", space="4", sz="12"):
    """Add a bottom border to a paragraph (used for section titles)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), space)
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_run_highlight(run, color_hex):
    """Set character shading (highlight) on a run."""
    rPr  = run._r.get_or_add_rPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color if (hex_color := color_hex) else color_hex)
    rPr.append(shd)

# ── Reusable formatters ────────────────────────────────────────────────────────
def body(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, space_after=5, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name  = FONT_BODY
    r.font.size  = Pt(size)
    r.italic     = italic
    r.font.color.rgb = BLACK
    return p

def section_heading(text, number=None):
    """Bold navy heading with a blue bottom border line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    if number:
        rn = p.add_run(f"{number}.  ")
        rn.font.name  = FONT_HEADING
        rn.font.size  = Pt(13)
        rn.bold       = True
        rn.font.color.rgb = ACCENT_BLUE
    r = p.add_run(text.upper())
    r.font.name  = FONT_HEADING
    r.font.size  = Pt(13)
    r.bold       = True
    r.font.color.rgb = DARK_BLUE
    set_para_border_bottom(p)
    return p

def mini_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name  = FONT_HEADING
    r.font.size  = Pt(11)
    r.bold       = True
    r.font.color.rgb = ACCENT_BLUE
    return p

def bullet_item(label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    rl = p.add_run(label + ": ")
    rl.bold          = True
    rl.font.name     = FONT_BODY
    rl.font.size     = Pt(11)
    rl.font.color.rgb = DARK_BLUE
    rb = p.add_run(text)
    rb.font.name     = FONT_BODY
    rb.font.size     = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE  (dark-navy banner style)
# ══════════════════════════════════════════════════════════════════════════════

# — top banner table (1 row, 1 col, dark blue fill) —
banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
bc = banner.rows[0].cells[0]
set_cell_bg(bc, "1A376C")
bc.width = Inches(6)

for para in bc.paragraphs:
    para.clear()

bp1 = bc.paragraphs[0]
bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp1.paragraph_format.space_before = Pt(22)
r = bp1.add_run("GESTALT LAWS")
r.bold           = True
r.font.name      = FONT_HEADING
r.font.size      = Pt(28)
r.font.color.rgb = WHITE

bp2 = bc.add_paragraph()
bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = bp2.add_run("OF PERCEPTUAL ORGANIZATION")
r2.bold           = True
r2.font.name      = FONT_HEADING
r2.font.size      = Pt(18)
r2.font.color.rgb = RGBColor(0xBD, 0xD7, 0xEE)

bp3 = bc.add_paragraph()
bp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp3.paragraph_format.space_before = Pt(10)
r3 = bp3.add_run("A Comprehensive Study in Gestalt Psychology")
r3.italic        = True
r3.font.name     = FONT_BODY
r3.font.size     = Pt(12)
r3.font.color.rgb = RGBColor(0xD6, 0xE4, 0xF7)

bp4 = bc.add_paragraph()
bp4.paragraph_format.space_after = Pt(22)

# — accent stripe (thin accent-blue table) —
stripe = doc.add_table(rows=1, cols=1)
stripe.style = "Table Grid"
sc = stripe.rows[0].cells[0]
set_cell_bg(sc, "2E74B5")
sc.paragraphs[0].paragraph_format.space_before = Pt(3)
sc.paragraphs[0].paragraph_format.space_after  = Pt(3)

doc.add_paragraph()

# — details card —
details_tbl = doc.add_table(rows=7, cols=2)
details_tbl.style = "Table Grid"
labels = ["Subject", "Topic", "Submitted By", "Roll No.", "Class / Section",
          "Submitted To", "Date"]
values = ["General Psychology",
          "Gestalt Laws of Perceptual Organization",
          "______________________ __",
          "________________________",
          "________________________",
          "________________________",
          "March 11, 2026"]

for i, (lbl, val) in enumerate(zip(labels, values)):
    lc = details_tbl.rows[i].cells[0]
    vc = details_tbl.rows[i].cells[1]
    set_cell_bg(lc, "1A376C")
    set_cell_bg(vc, "EBF3FB")

    lc.paragraphs[0].clear()
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(4)
    lr = lp.add_run(lbl.upper())
    lr.bold           = True
    lr.font.name      = FONT_HEADING
    lr.font.size      = Pt(10)
    lr.font.color.rgb = WHITE

    vc.paragraphs[0].clear()
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(4)
    vp.paragraph_format.space_after  = Pt(4)
    vr = vp.add_run(val)
    vr.font.name      = FONT_BODY
    vr.font.size      = Pt(10)
    vr.font.color.rgb = DARK_BLUE

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Abstract")
body(
    "This assignment examines the Gestalt Laws of Perceptual Organization — a set of "
    "principles formulated by German psychologists Koffka, Köhler, and Wertheimer to explain "
    "how the human mind structures sensory data into coherent wholes. The paper discusses "
    "the theoretical background of Gestalt psychology, reviews each major perceptual law "
    "with illustrative examples, and evaluates their relevance to contemporary fields such "
    "as design, education, and clinical psychology. The central thesis — that the whole is "
    "greater than the sum of its parts — is explored across cognitive, developmental, and "
    "applied dimensions.",
    italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Background: The Gestalt Movement")
body(
    "Gestalt psychology emerged in Germany in the early twentieth century as a direct "
    "reaction against the structuralist approach of Wilhelm Wundt, which sought to reduce "
    "mental experience into elementary sensory components. The founders — Max Wertheimer "
    "(1880–1943), Kurt Koffka (1886–1941), and Wolfgang Köhler (1887–1967) — argued that "
    "such reductionism misrepresented how the mind actually works. The word 'Gestalt' "
    "(German: form, shape, whole) encapsulates their belief that perception is inherently "
    "holistic: the mind does not passively receive isolated bits of data but actively "
    "organises them into structured, meaningful patterns."
)
body(
    "Wertheimer's 1912 experiments on apparent motion (the phi phenomenon) demonstrated "
    "that a sequence of still images could produce the perception of motion — a phenomenon "
    "irreducible to its individual frames. This became the empirical cornerstone of Gestalt "
    "theory. Koffka later systematised the perceptual laws in 'Principles of Gestalt "
    "Psychology' (1935), and Köhler contributed insights from animal learning that "
    "reinforced the holistic view of cognition."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PERCEPTION: FIGURE AND GROUND
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Perception: Figure and Ground", number=1)
body(
    "The most fundamental distinction in perceptual organisation is between figure and "
    "ground. In any visual scene, certain elements stand out as distinct objects (the figure) "
    "while others recede as a background context (the ground). The figure is perceived as "
    "closer, better defined, and more memorable; the ground is perceived as further away, "
    "less defined, and continuous."
)
body(
    "This distinction is not fixed and can fluctuate depending on attention and context. "
    "The classic Rubin's Vase demonstrates this reversibility: the same image can be "
    "perceived either as a white vase (figure) on a black ground, or as two black faces "
    "(figure) on a white ground. The oscillation highlights that perception is constructive "
    "— the brain imposes interpretation rather than passively recording input."
)

# callout box via single-cell table
cbox = doc.add_table(rows=1, cols=1)
cbox.style = "Table Grid"
cc = cbox.rows[0].cells[0]
set_cell_bg(cc, "EBF3FB")
cp = cc.paragraphs[0]
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after  = Pt(6)
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run(
    '"Just as the focus of consciousness fluctuates, in the same way various stimuli '
    'keep oscillating between figure and ground."'
)
cr.italic        = True
cr.font.name     = FONT_BODY
cr.font.size     = Pt(11)
cr.font.color.rgb = DARK_BLUE
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4–9. THE GESTALT LAWS (detail table cards)
# ══════════════════════════════════════════════════════════════════════════════
laws = [
    (2, "Law of Proximity",
     "Definition",
     "Elements that are spatially close to one another tend to be grouped together "
     "and perceived as a single unit.",
     "Example / Application",
     "In typography, words (closely spaced letters) are read as units rather than "
     "individual characters. In UI design, related buttons are placed near each other "
     "so users understand they share a function. Even when proximity conflicts with "
     "similarity, proximity often wins as the dominant grouping cue."),

    (3, "Law of Similarity",
     "Definition",
     "Elements that share visual properties — shape, size, colour, texture, or "
     "orientation — are perceived as belonging to the same group.",
     "Example / Application",
     "A grid of alternating circles and squares is perceived as rows of circles and "
     "rows of squares. Marketers exploit similarity by giving a product line consistent "
     "packaging so consumers mentally group the products together."),

    (4, "Law of Continuity",
     "Definition",
     "The perceptual system prefers smooth, continuous paths over abrupt changes in "
     "direction. Elements arranged along a smooth curve or line are perceived as "
     "belonging together.",
     "Example / Application",
     "When two curved lines cross, we perceive two intact crossing curves rather than "
     "four separate segments joining at a point. Road designers and animators use "
     "continuity to create natural-feeling trajectories."),

    (5, "Law of Closure",
     "Definition",
     "The mind fills in missing information to perceive incomplete shapes as complete "
     "figures. Perception 'closes' gaps automatically.",
     "Example / Application",
     "A circle drawn with a small gap is still perceived as a circle. The WWF panda "
     "logo uses closure — the outline is incomplete yet perceived as a whole panda. "
     "Closure enables object recognition even under partial occlusion."),

    (6, "Law of Symmetry",
     "Definition",
     "Symmetrical elements are grouped together and perceived as a single cohesive "
     "figure, even across empty space, because symmetry implies a common origin.",
     "Example / Application",
     "The brackets in '[  ]' are perceived as enclosing one region rather than two "
     "independent marks. Symmetry contributes to the aesthetic quality of logos, "
     "architecture, and faces."),

    (7, "Law of Prägnanz (Good Form)",
     "Definition",
     "The overarching Gestalt law: the mind always resolves ambiguous stimuli into "
     "the simplest, most regular, and most stable interpretation possible. "
     "'Prägnanz' (German: precision) captures this drive toward economy.",
     "Example / Application",
     "An Olympic-rings-style arrangement of overlapping circles is perceived as "
     "complete overlapping circles, not as a complex irregular polygon. Minimalist "
     "corporate logos leverage Prägnanz for instant recognition."),

    (8, "Law of Common Fate",
     "Definition",
     "Elements that move together in the same direction and at the same speed are "
     "perceived as a single group, regardless of their visual similarity.",
     "Example / Application",
     "A flock of birds flying in formation is perceived as a single unit. In data "
     "visualisation, animated particles moving together signal membership in the same "
     "data cluster. Common fate extends grouping into the time dimension."),
]

for num, title, lbl1, txt1, lbl2, txt2 in laws:
    section_heading(title, number=num)
    # two-column info card
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"

    # header row
    h0 = tbl.rows[0].cells[0]
    h1 = tbl.rows[0].cells[1]
    set_cell_bg(h0, "2E74B5")
    set_cell_bg(h1, "2E74B5")
    for cell, text in ((h0, lbl1), (h1, lbl2)):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text.upper())
        r.bold           = True
        r.font.name      = FONT_HEADING
        r.font.size      = Pt(10)
        r.font.color.rgb = WHITE

    # content row
    c0 = tbl.rows[1].cells[0]
    c1 = tbl.rows[1].cells[1]
    set_cell_bg(c0, "F2F9FF")
    set_cell_bg(c1, "FAFAFA")
    for cell, text in ((c0, txt1), (c1, txt2)):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(text)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(10)
        r.font.color.rgb = BLACK

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. COMPARATIVE TABLE OF ALL LAWS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Summary Comparison of Gestalt Laws")
body("The table below provides a quick-reference overview of all eight Gestalt laws.")

sum_tbl = doc.add_table(rows=1, cols=4)
sum_tbl.style = "Table Grid"
col_headers = ["#", "Law", "Core Idea", "Key Domain"]
for i, h in enumerate(col_headers):
    cell = sum_tbl.rows[0].cells[i]
    set_cell_bg(cell, "1A376C")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(h.upper())
    r.bold           = True
    r.font.name      = FONT_HEADING
    r.font.size      = Pt(10)
    r.font.color.rgb = WHITE

summary_data = [
    ("1", "Figure & Ground",  "Objects vs. background",         "Attention & awareness"),
    ("2", "Proximity",        "Nearness implies grouping",      "Layout & navigation"),
    ("3", "Similarity",       "Likeness implies grouping",      "Pattern recognition"),
    ("4", "Continuity",       "Smooth paths are preferred",     "Motion & animation"),
    ("5", "Closure",          "Incomplete figures completed",   "Logo & icon design"),
    ("6", "Symmetry",         "Balanced forms unified",         "Aesthetics & branding"),
    ("7", "Prägnanz",         "Simplest form always chosen",    "Minimalist design"),
    ("8", "Common Fate",      "Co-moving elements grouped",     "Data visualisation"),
]
alt_colors = ["EBF3FB", "FAFAFA"]
for idx, row_d in enumerate(summary_data):
    row_cells = sum_tbl.add_row().cells
    bg = alt_colors[idx % 2]
    for j, val in enumerate(row_d):
        set_cell_bg(row_cells[j], bg)
        p = row_cells[j].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        r.font.name      = FONT_BODY
        r.font.size      = Pt(10)
        r.font.color.rgb = DARK_BLUE if j == 1 else BLACK
        if j == 1:
            r.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. CRITICAL EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Critical Evaluation")

mini_heading("Strengths")
body(
    "Gestalt laws are supported by a wealth of empirical evidence from psychophysical "
    "experiments and neuroimaging studies. They offer parsimonious explanations for "
    "a wide variety of perceptual phenomena and provide actionable design guidelines. "
    "The laws have proven robust across cultures and age groups, suggesting they reflect "
    "deep-seated neural mechanisms rather than learned conventions."
)

mini_heading("Limitations")
body(
    "Critics note that Gestalt theory is largely descriptive rather than mechanistic — it "
    "identifies what the brain does without fully explaining how. The laws can also be "
    "vague and sometimes make conflicting predictions when multiple laws apply simultaneously. "
    "Furthermore, early Gestalt researchers relied on introspective and qualitative "
    "methods that do not meet modern standards of experimental rigor."
)

mini_heading("Modern Neuroscience Perspective")
body(
    "Contemporary cognitive neuroscience has begun to uncover the neural substrates of "
    "Gestalt grouping. Research in primary visual cortex (V1) and higher areas (V4, LOC) "
    "shows that neurons respond differentially to grouped vs. ungrouped stimuli, providing "
    "biological grounding for Gestalt phenomena. Predictive coding frameworks suggest "
    "Gestalt laws reflect the brain's probabilistic model of a world filled with "
    "continuous, symmetric, and proximate objects."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. APPLICATIONS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Real-World Applications")
body("Gestalt principles permeate numerous professional and academic domains:")

app_items = [
    ("Graphic & UI Design",
     "Proximity and similarity guide information hierarchy; closure and figure-ground create "
     "memorable logos (e.g., FedEx hidden arrow, NBC peacock)."),
    ("Architecture & Urban Planning",
     "Symmetry and Prägnanz inform façade composition and public space design so that "
     "environments feel ordered and navigable."),
    ("Data Visualisation",
     "Common fate and proximity determine how viewers group data points into trends and "
     "clusters in charts and dashboards."),
    ("Education",
     "Teachers use proximity (grouping related content on a slide) and similarity "
     "(consistent formatting for related concepts) to reduce cognitive load."),
    ("Clinical Neuropsychology",
     "Disruptions in Gestalt perception (e.g., inability to perceive figures from ground) "
     "can be diagnostic indicators of conditions such as schizophrenia or prosopagnosia."),
    ("Film & Animation",
     "Continuity and common fate underpin scene editing and character motion — the "
     "audience perceptually binds frames into coherent narrative sequences."),
]
for lbl, txt in app_items:
    bullet_item(lbl, txt)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("Conclusion")
body(
    "Gestalt laws of perceptual organisation represent one of psychology's most enduring "
    "contributions to understanding the human mind. By demonstrating that organised wholes "
    "cannot be reduced to the sum of their parts, Wertheimer, Koffka, and Köhler "
    "fundamentally reshaped the study of perception and laid groundwork that still "
    "influences cognitive science, neuroscience, and design today. The principles of "
    "proximity, similarity, continuity, closure, symmetry, Prägnanz, common fate, and "
    "figure-ground are not merely academic curiosities — they are the invisible grammar "
    "through which the human visual system makes sense of the world."
)
body(
    "As artificial intelligence systems increasingly attempt to replicate human perception, "
    "Gestalt theory offers a principled benchmark: a truly perceptive system must organise "
    "input into meaningful wholes rather than processing isolated features in isolation. "
    "The relevance of Gestalt psychology, therefore, extends well into the twenty-first century."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
section_heading("References")

refs = [
    ("Wertheimer, M.", "1923",
     "Laws of organization in perceptual forms. In W. D. Ellis (Ed.), "
     "A source book of Gestalt psychology. Harcourt, Brace."),
    ("Koffka, K.", "1935",
     "Principles of Gestalt psychology. Harcourt, Brace."),
    ("Köhler, W.", "1947",
     "Gestalt psychology: An introduction to new concepts in modern psychology. Liveright."),
    ("Goldstein, E. B.", "2019",
     "Sensation and perception (10th ed.). Cengage Learning."),
    ("Morgan, C. T., King, R. A., Weisz, J. R., & Schopler, J.", "2012",
     "Introduction to psychology (7th ed.). Tata McGraw-Hill."),
    ("Wagemans, J. et al.", "2012",
     "A century of Gestalt psychology in visual perception. "
     "Psychological Bulletin, 138(6), 1172–1217."),
]

for author, year, detail in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after   = Pt(5)
    ra = p.add_run(author + " ")
    ra.bold          = True
    ra.font.name     = FONT_BODY
    ra.font.size     = Pt(10)
    ra.font.color.rgb = DARK_BLUE
    ry = p.add_run(f"({year}). ")
    ry.font.name     = FONT_BODY
    ry.font.size     = Pt(10)
    rd = p.add_run(detail)
    rd.font.name     = FONT_BODY
    rd.font.size     = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER — page numbers
# ══════════════════════════════════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Gestalt Laws of Perceptual Organization  |  General Psychology  |  Page ")
    fr.font.name     = FONT_BODY
    fr.font.size     = Pt(9)
    fr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    # field for auto page number
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_xml = fp.add_run()._r
    run_xml.append(fldChar1)
    run_xml.append(instrText)
    run_xml.append(fldChar2)

# ══════════════════════════════════════════════════════════════════════════════
output_path = r"c:\Users\Abdul\Desktop\Assignment\Gestalt_Laws_Assignment_v2.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
