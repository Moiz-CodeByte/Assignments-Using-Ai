from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helper: set font ──────────────────────────────────────────────────────────
def fmt(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, center=False, underline=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p

def add_body(doc, text, justify=True, spacing_after=6):
    p = doc.add_paragraph(text)
    p.style.font.name = "Times New Roman"
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(spacing_after)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p

def add_subheading(doc, text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("GESTALT LAWS OF PERCEPTUAL ORGANIZATION")
r.bold = True
r.underline = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("An Assignment on Perception and Gestalt Psychology")
r2.font.size = Pt(13)
r2.font.name = "Times New Roman"
r2.italic = True

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ("Subject:", "General Psychology"),
    ("Topic:", "Gestalt Laws of Perceptual Organization"),
    ("Submitted by:", "________________________"),
    ("Roll No.:", "________________________"),
    ("Class:", "________________________"),
    ("Submitted to:", "________________________"),
    ("Date:", "March 11, 2026"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = p.add_run(f"{label}  ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(12)
    r_lbl.font.name = "Times New Roman"
    r_val = p.add_run(value)
    r_val.font.size = Pt(12)
    r_val.font.name = "Times New Roman"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "INTRODUCTION", size=14, underline=True)
doc.add_paragraph()

add_body(doc,
    "Perception is the process through which individuals interpret and organise sensory "
    "information to give meaning to the environment. In perception, we view figure in "
    "isolation from ground. The individual is influenced by a number of stimuli in the "
    "environment. All the things in the environment cannot be brought into the field of "
    "consciousness. Some things become clear because of the individual's interest whereas "
    "the rest remain vague and indistinct. The thing that is clear and conspicuous in the "
    "environment is called figure. To put it differently, the thing that occupies an important "
    "position in the environment becomes figure. The things which are vague, indistinct and "
    "unimportant recede into the background and are referred to as ground. Figure and ground "
    "are analogous to field and center of consciousness: Ground is the field of consciousness "
    "whereas figure is the center of consciousness. Just as the focus of consciousness "
    "fluctuates, in the same way various stimuli keep oscillating between figure and ground."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FACTORS INFLUENCING PERCEPTION  (table referenced in image 2)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "FACTORS INFLUENCING PERCEPTION", size=13, underline=True)
doc.add_paragraph()

add_body(doc,
    "Perception is not a passive recording of stimuli; it is influenced by both external "
    "(stimulus) factors and internal (personal) factors. The major factors are summarised below:"
)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
headers = ["Stimulus Factors", "Personal Factors", "Socio-Cultural Factors"]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for run in hdr_cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

rows_data = [
    ("Size",          "Motivation",           "Social Values"),
    ("Intensity",     "Interest",             "Cultural background"),
    ("Contrast",      "Emotional Condition",  "Social Values"),
    ("Colour",        "Beliefs",              "Suggestion"),
    ("Loudness",      "Personal qualities",   ""),
    ("",              "Physical & mental health", ""),
]
for row_d in rows_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_d):
        row_cells[i].text = val
        for run in row_cells[i].paragraphs[0].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# GESTALT LAWS — main heading
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "GESTALT LAWS OF PERCEPTUAL ORGANIZATION", size=14, underline=True)
doc.add_paragraph()

add_body(doc,
    "Gestalt is the school of thought that laid emphasis on studying things and factors as a "
    "whole. German psychologists Koffka, Wolfgang Kohler and Max Wertheimer are its main "
    "exponents. Gestalt psychologists explained the laws of perceptual organization. Hence "
    "these laws are referred to as Gestalt laws of perceptual organization."
)

add_body(doc,
    "These psychologists maintain that human perceptual processes reflect mental organization. "
    "They assert that \"a whole is greater than the sum of its parts\", meaning the complete "
    "perception of a thing is more perfect and organized than the analysis of its constituents. "
    "They say that the individual organises things or stimuli into groups during perception. "
    "He views the various stimuli in the visual area as organized wholes. The laws whereby "
    "stimuli are grouped are called laws of perceptual organization. Their explanation is "
    "given as under:"
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 1 – FIGURE AND GROUND
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "1. FIGURE AND GROUND")
add_body(doc,
    "The most fundamental Gestalt principle is the distinction between figure and ground. "
    "When we look at a scene, we naturally separate objects (figure) from the background "
    "(ground). The figure appears more defined, closer, and stands out; the ground appears "
    "less defined and recedes. This relationship is not fixed — what is figure at one moment "
    "can become ground at another. A classic demonstration is the Rubin's Vase illusion, "
    "where the viewer alternates between seeing a vase (figure on a dark ground) and two "
    "faces (figure on a light ground). This oscillation illustrates that perception is an "
    "active, constructive process shaped by attention and context."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 2 – PROXIMITY
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "2. LAW OF PROXIMITY")
add_body(doc,
    "The law of proximity states that objects or stimuli that are close to one another tend "
    "to be grouped together and perceived as a single unit or pattern. When elements are "
    "near each other, the perceptual system automatically clusters them. For example, a set "
    "of dots arranged in pairs will be perceived as groups of two rather than as individual "
    "dots. This principle is widely applied in design, layout, and visual communication — "
    "items placed close together are understood as related. Proximity is one of the most "
    "powerful grouping cues and operates even when the individual elements are dissimilar."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 3 – SIMILARITY
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "3. LAW OF SIMILARITY")
add_body(doc,
    "According to the law of similarity, elements that are similar to each other in shape, "
    "size, colour, texture, or orientation tend to be grouped together. The perceptual system "
    "links similar items and treats them as belonging to the same object or category. For "
    "instance, in a grid containing circles and squares, the circles will be perceived as "
    "one group and the squares as another. Similarity acts as a powerful organising principle "
    "even when elements are not in close proximity to one another. This law explains why "
    "we see patterns such as rows and columns in uniform grids."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 4 – CONTINUITY
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "4. LAW OF CONTINUITY (Good Continuation)")
add_body(doc,
    "The law of continuity holds that the human perceptual system tends to perceive elements "
    "as belonging to a continuous, flowing line or curve rather than as abrupt or discontinuous "
    "segments. When lines or curves intersect, we perceive them as continuing along their "
    "original path. For example, two crossing curved lines are seen as two intact curves rather "
    "than four separate segments meeting at a point. This Gestalt principle helps explain how "
    "we follow the trajectory of moving objects, read connected handwriting, and interpret "
    "overlapping forms. It reflects the mind's preference for smooth, regular, and predictable "
    "patterns over jagged or irregular ones."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 5 – CLOSURE
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "5. LAW OF CLOSURE")
add_body(doc,
    "The law of closure refers to the tendency of the perceptual system to complete incomplete "
    "figures or forms. When we encounter an incomplete shape — a circle with a gap, or a "
    "partially visible object — the mind fills in the missing information to perceive a whole, "
    "complete figure. This principle enables us to recognise objects even when parts of them "
    "are hidden or occluded. For example, a triangle drawn with broken lines is still perceived "
    "as a complete triangle. Closure demonstrates that perception is not merely the recording "
    "of sensory data but an active process of inference and completion guided by expectations "
    "and prior knowledge."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 6 – SYMMETRY
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "6. LAW OF SYMMETRY")
add_body(doc,
    "The law of symmetry states that the mind perceives objects as symmetrical and structured "
    "around a central point or axis. Symmetrical elements tend to be grouped together and "
    "perceived as a single unified figure, even when they are separated by some distance. "
    "This principle promotes perceptual stability and simplicity. For example, two symmetrical "
    "bracket shapes facing each other — such as '[ ]' — are perceived as enclosing a single "
    "region rather than as two independent shapes. Symmetry contributes to the aesthetic "
    "appeal of visual designs and plays a role in recognising faces and biological forms."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 7 – PRÄGNANZ (Law of Good Form / Simplicity)
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "7. LAW OF PRÄGNANZ (Law of Good Form / Simplicity)")
add_body(doc,
    "The law of Prägnanz — the overarching Gestalt principle — states that the perceptual "
    "system always organises stimuli in the simplest, most regular, and most stable form "
    "possible. \"Prägnanz\" is a German word meaning precision or conciseness. When presented "
    "with ambiguous or complex visual information, the brain automatically resolves it into "
    "the simplest interpretation. For example, an overlapping arrangement of circles is "
    "perceived as complete overlapping circles rather than as a complex irregular shape. "
    "Prägnanz is the fundamental drive behind all other Gestalt laws — the mind seeks "
    "order, balance, and simplicity in every perceptual experience."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# LAW 8 – COMMON FATE
# ══════════════════════════════════════════════════════════════════════════════
add_subheading(doc, "8. LAW OF COMMON FATE")
add_body(doc,
    "The law of common fate states that elements that move in the same direction or at the "
    "same rate tend to be perceived as belonging to the same group or as forming a single "
    "unit. This principle extends Gestalt grouping into the temporal dimension of motion. "
    "For instance, birds flying in the same direction are perceived as a single flock. "
    "Similarly, particles drifting in the same direction are grouped perceptually. Common fate "
    "is fundamental in perceiving biological motion, tracking objects in motion, and "
    "understanding dynamic visual scenes."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SIGNIFICANCE / APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "SIGNIFICANCE AND APPLICATIONS OF GESTALT LAWS", size=13, underline=True)
doc.add_paragraph()

add_body(doc,
    "Gestalt laws have profound implications across multiple fields:"
)

applications = [
    ("Graphic Design and Art:",
     "Designers use proximity, similarity, and closure to create clear, aesthetically "
     "pleasing layouts. Logos often exploit these principles (e.g., the FedEx arrow "
     "hidden between letters uses closure and figure-ground)."),

    ("User Interface (UI) Design:",
     "Web and app interfaces group related controls together (proximity) and use "
     "consistent styling (similarity) so users intuitively understand the interface."),

    ("Education and Teaching:",
     "Teachers can organise information on boards or slides using Gestalt principles "
     "to direct student attention and aid memory."),

    ("Advertising and Marketing:",
     "Advertisements structure visual layouts to make the brand or product the figure "
     "while pushing competing information to the ground."),

    ("Clinical Psychology:",
     "Understanding perceptual organisation helps clinicians with assessments such as "
     "the Rorschach inkblot test, which exploits figure-ground ambiguity."),
]

for title_text, body_text in applications:
    p = doc.add_paragraph(style="List Bullet")
    r_t = p.add_run(title_text + "  ")
    r_t.bold = True
    r_t.font.name = "Times New Roman"
    r_t.font.size = Pt(12)
    r_b = p.add_run(body_text)
    r_b.font.name = "Times New Roman"
    r_b.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "CONCLUSION", size=13, underline=True)
doc.add_paragraph()

add_body(doc,
    "Gestalt laws of perceptual organisation provide a powerful framework for understanding "
    "how the human mind structures sensory input into meaningful wholes. Rather than perceiving "
    "the world as a collection of isolated stimuli, the perceptual system imposes order, "
    "continuity, and coherence. The principles of figure-ground, proximity, similarity, "
    "continuity, closure, symmetry, Prägnanz, and common fate collectively demonstrate that "
    "perception is an active, constructive process. The assertion that 'the whole is greater "
    "than the sum of its parts' remains as relevant today as when Wertheimer, Koffka, and "
    "Köhler first articulated it. A thorough understanding of these laws is indispensable "
    "not only in psychology but in every domain that relies on visual communication and "
    "human-centred design."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "REFERENCES", size=13, underline=True)
doc.add_paragraph()

refs = [
    "Wertheimer, M. (1923). Laws of organization in perceptual forms. In W. D. Ellis (Ed.), "
    "A source book of Gestalt psychology. Harcourt, Brace.",

    "Koffka, K. (1935). Principles of Gestalt psychology. Harcourt, Brace.",

    "Köhler, W. (1947). Gestalt psychology: An introduction to new concepts in modern "
    "psychology. Liveright.",

    "Goldstein, E. B. (2019). Sensation and perception (10th ed.). Cengage Learning.",

    "Morgan, C. T., King, R. A., Weisz, J. R., & Schopler, J. (2012). Introduction to "
    "psychology (7th ed.). Tata McGraw-Hill.",

    "Course notes / reference material provided by the instructor.",
]

for ref in refs:
    p = doc.add_paragraph(ref, style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"c:\Users\Abdul\Desktop\Assignment\Gestalt_Laws_Assignment.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
