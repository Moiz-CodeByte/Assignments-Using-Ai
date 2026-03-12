from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margins
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1.25)

NAVY = RGBColor(0x1F, 0x39, 0x64)

def h(text, size=13, underline=False, center=False, space_before=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.underline = underline
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY

def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)
    return p

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
h("Gestalt Laws of Perceptual Organization", size=18, center=True, space_before=0)
h("Assignment — General Psychology", size=12, center=True, space_before=6)
doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=2)
tbl.style = "Table Grid"
info = [("Submitted By", "________________________"),
        ("Roll No.",     "________________________"),
        ("Class",        "________________________"),
        ("Submitted To", "________________________"),
        ("Date",         "March 11, 2026")]
for i, (l, v) in enumerate(info):
    lc, vc = tbl.rows[i].cells
    lp = lc.paragraphs[0]
    lr = lp.add_run(l)
    lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11)
    vp = vc.paragraphs[0]
    vr = vp.add_run(v)
    vr.font.name = "Arial"; vr.font.size = Pt(11)

doc.add_page_break()

# ── INTRODUCTION ──────────────────────────────────────────────────────────────
h("Introduction", underline=True)
body(
    "Gestalt psychology is a school of thought founded by German psychologists Koffka, "
    "Köhler, and Wertheimer. It studies how the mind organizes sensory information into "
    "meaningful wholes. Their core belief is that \"the whole is greater than the sum of "
    "its parts.\" The laws they proposed explain how we group and interpret visual stimuli."
)

# ── LAWS ──────────────────────────────────────────────────────────────────────
laws = [
    ("1. Figure and Ground",
     "We naturally separate a scene into a main object (figure) and a background (ground). "
     "The figure appears closer and well-defined; the ground appears distant and vague. "
     "This relationship can flip — e.g., Rubin's Vase appears as either a vase or two faces."),

    ("2. Proximity",
     "Objects that are close to each other are perceived as a group. For example, dots "
     "placed in pairs look like groups of two rather than individual dots."),

    ("3. Similarity",
     "Objects that look alike — same shape, color, or size — are grouped together. "
     "A grid of circles and squares is seen as separate groups of each shape."),

    ("4. Continuity",
     "The eye follows smooth continuous lines rather than abrupt changes. Two crossing "
     "curves are seen as two intact curves, not as four separate pieces."),

    ("5. Closure",
     "The mind fills in gaps to perceive a complete shape. A broken circle is still "
     "seen as a circle. This allows us to recognize objects even when partially hidden."),

    ("6. Symmetry",
     "Symmetrical elements are grouped and seen as one unified figure. Balanced forms "
     "feel stable and complete, even when separated by space."),

    ("7. Prägnanz (Good Form)",
     "The mind always chooses the simplest possible interpretation of a stimulus. "
     "Ambiguous images are resolved into the most regular, orderly form available."),

    ("8. Common Fate",
     "Elements moving in the same direction at the same speed are grouped together. "
     "A flock of birds flying in formation is perceived as a single unit."),
]

for title, text in laws:
    h(title, size=12, underline=False, space_before=8)
    body(text)

# ── CONCLUSION ────────────────────────────────────────────────────────────────
h("Conclusion", underline=True)
body(
    "Gestalt laws show that perception is an active process — the brain organizes "
    "sensory input into structured, meaningful patterns. These principles are widely "
    "applied in graphic design, user interfaces, education, and psychology to guide "
    "how information is presented and understood."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
h("References", underline=True)
refs = [
    "Koffka, K. (1935). Principles of Gestalt Psychology. Harcourt, Brace.",
    "Köhler, W. (1947). Gestalt Psychology. Liveright.",
    "Wertheimer, M. (1923). Laws of organization in perceptual forms.",
    "Goldstein, E. B. (2019). Sensation and Perception (10th ed.). Cengage.",
]
for ref in refs:
    p = doc.add_paragraph(ref, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10)

out = r"c:\Users\Abdul\Desktop\Assignment\Gestalt_Laws_Assignment_v3.docx"
doc.save(out)
print(f"Saved: {out}")
