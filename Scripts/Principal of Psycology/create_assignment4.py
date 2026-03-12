from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margins
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1.25)

NAVY = RGBColor(0x1F, 0x39, 0x64)

def h(text, size=13, underline=False, center=False, space_before=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.underline = underline
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY

def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)
    return p

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
h("How to Help Employees Achieve Work-Life Balance", size=18, center=True, space_before=0)
h("Assignment — Professional Practices in Management", size=12, center=True, space_before=6)
doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=2)
tbl.style = "Table Grid"
info = [("Submitted By", "________________________"),
        ("Roll No.",     "________________________"),
        ("Class",        "________________________"),
        ("Submitted To", "________________________"),
        ("Date",         "March 13, 2026")]
for i, (l, v) in enumerate(info):
    lc, vc = tbl.rows[i].cells
    lp = lc.paragraphs[0]
    lr = lp.add_run(l)
    lr.bold = True; lr.font.name = "Arial"; lr.font.size = Pt(11)
    vp = vc.paragraphs[0]
    vr = vp.add_run(v)
    vr.font.name = "Arial"; vr.font.size = Pt(11)

doc.add_page_break()

# ── INTRODUCTION ──────────────────────────────────────────────────────────────
h("Introduction", underline=True)
body(
    "Work-life balance refers to the equilibrium between time spent on career and "
    "professional responsibilities and time dedicated to personal life, family, health, "
    "and leisure. In today's fast-paced corporate environment, maintaining this balance "
    "is increasingly challenging yet critical for employee well-being, productivity, and "
    "organizational success. Companies that prioritize work-life balance benefit from "
    "reduced turnover, improved morale, and increased employee satisfaction."
)

# ── STRATEGIES ────────────────────────────────────────────────────────────────
h("Key Strategies for Achieving Work-Life Balance", underline=True, space_before=10)

strategies = [
    ("1. Implement Flexible Work Arrangements",
     "Offer flexible schedules, remote work options, and compressed work weeks. Allowing "
     "employees to adjust their work hours to fit personal commitments reduces stress and "
     "increases autonomy. This flexibility enables parents to manage childcare, students "
     "to pursue education, and individuals to maintain personal wellness."),

    ("2. Set Clear Work Boundaries",
     "Establish policies that respect off-hours time. Discourage after-hours emails, "
     "excessive overtime, and weekend work unless necessary. Clear boundaries help employees "
     "mentally disconnect from work and recharge. Communication that 'work can wait' is essential."),

    ("3. Promote Health and Wellness Programs",
     "Provide gym memberships, mental health counseling, stress management workshops, and "
     "wellness initiatives. Active wellness programs show that the organization values "
     "employee health. These programs reduce burnout and improve overall quality of life."),

    ("4. Encourage Time Off",
     "Ensure employees use their vacation and paid time off. Leaders should model this behavior "
     "by taking their own leave. Regular breaks prevent burnout and allow for personal rejuvenation. "
     "Consider unlimited PTO or generous leave policies to demonstrate trust."),

    ("5. Optimize Workload and Resource Management",
     "Ensure workloads are realistic and resources are adequate. Overwork is a primary cause of "
     "poor work-life balance. Regular workload assessments and fair distribution prevent bottlenecks "
     "and allow employees to complete tasks during work hours."),

    ("6. Foster a Supportive Company Culture",
     "Build a culture where work-life balance is respected, not penalized. Employees who fear "
     "career consequences for leaving on time will overwork. Leadership must communicate that "
     "balance is valued and supported at all organizational levels."),

    ("7. Provide Career Development Opportunities",
     "Invest in employee growth through training, mentorship, and clear career paths. Job security "
     "and clear advancement opportunities reduce anxiety and allow employees to plan their lives "
     "with confidence."),

    ("8. Use Technology Wisely",
     "While technology enables distant work, it can also blur boundaries. Establish protocols for "
     "communication technology use. Implement 'do not disturb' hours and encourage downtime from screens "
     "to prevent digital burnout."),

    ("9. Support Working Parents",
     "Offer childcare assistance, parental leave, and flexible schedules for parents. Recognize that "
     "parenting is demanding and provide targeted support. This demonstrates inclusivity and reduces "
     "stress for a significant portion of the workforce."),

    ("10. Lead by Example",
     "Managers and executives must model healthy work-life balance. When leaders consistently work "
     "late or skip vacations, it signals that balance is not truly valued. Leadership behavior sets "
     "the organizational tone."),
]

for title, text in strategies:
    h(title, size=12, underline=False, space_before=8)
    body(text)

# ── BENEFITS ──────────────────────────────────────────────────────────────────
h("Benefits of Work-Life Balance", underline=True, space_before=10)

benefits = [
    ("For Employees",
     "Reduced stress and anxiety, better physical and mental health, improved relationships, "
     "increased job satisfaction, higher self-esteem, and greater overall life happiness."),

    ("For Organizations",
     "Lower absenteeism and turnover rates, higher productivity and engagement, improved "
     "employee retention, reduced healthcare costs, enhanced company reputation, and stronger "
     "employee loyalty."),
]

for title, text in benefits:
    h(title, size=12, underline=False, space_before=8)
    body(text)

# ── CHALLENGES ────────────────────────────────────────────────────────────────
h("Common Challenges", underline=True, space_before=10)
challenges = [
    "Heavy workloads and unrealistic deadlines",
    "Always-on work culture and constant connectivity",
    "Fear of career penalties for taking time off",
    "Inadequate staffing and resource limitations",
    "Lack of management understanding and support",
    "Global teams across time zones creating pressure",
]
for challenge in challenges:
    p = doc.add_paragraph(challenge, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)

# ── CONCLUSION ────────────────────────────────────────────────────────────────
h("Conclusion", underline=True, space_before=10)
body(
    "Work-life balance is not a luxury but a necessity for employee well-being and "
    "organizational success. By implementing flexible arrangements, setting boundaries, "
    "promoting wellness, and fostering a supportive culture, companies can create an "
    "environment where employees thrive both professionally and personally. Organizations "
    "that invest in work-life balance initiatives build stronger, more engaged, and more "
    "productive teams that drive long-term success."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
h("References", underline=True, space_before=10)
refs = [
    "American Psychological Association. (2023). Work and Well-being Survey.",
    "Friedman, S. D., & Greenhaus, J. H. (2000). Work and Family - Allies or Enemies?",
    "Kelliher, C., Richardson, J., & Boiarintseva, G. (2019). Have you got it all? The relationship between flexibility and gender in professional services.",
    "Society for Human Resource Management. (2023). Employee Benefits Survey.",
    "Williams, J. C., Blair-Loy, M., & Berdahl, J. L. (2013). Cultural schemas, social class, and the flexibility stigma.",
]
for ref in refs:
    p = doc.add_paragraph(ref, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10)

out = r"t:\Abdul\Assignments Using ai\Word Files\Work_Life_Balance_Assignment_v4.docx"
doc.save(out)
print(f"Saved: {out}")
